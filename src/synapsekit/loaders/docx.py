from __future__ import annotations

import os

from .base import Document


class DocxLoader:
    """Load text from a Microsoft Word (.docx) file.

    Requires ``python-docx``::

        pip install synapsekit[docx]
    """

    def __init__(self, path: str) -> None:
        self._path = path

    def load(self) -> list[Document]:
        if not os.path.exists(self._path):
            raise FileNotFoundError(f"File not found: {self._path}")

        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("python-docx required: pip install synapsekit[docx]") from None

        doc = DocxDocument(self._path)
        text = "\n".join(p.text for p in doc.paragraphs if p.text)
        return [Document(text=text, metadata={"source": self._path})]
